import tkinter as tk
from tkinter import ttk, messagebox
from docx import Document
import requests
import json

# Function to generate the Word document
def generate_docx(fields):
    # Create a new Word document
    doc = Document()

    # Add a title
    doc.add_heading('AI Text Generator', level=1)
    doc.add_heading('', level=2)

    # Add a table with three columns: Field Name, Description, and Comments
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'

    # Set column headers
    table.cell(0, 0).text = 'Field Name'
    table.cell(0, 1).text = 'Description'
    table.cell(0, 2).text = 'Comments'

    # Add a row for each field
    for i in range(len(fields) + 1):  # Loop over len(fields) + 1
        if i < len(fields):
            field = fields[i]
            row = table.add_row().cells
            row[0].text = field['name']
            row[1].text = field['description']
            row[2].text = field['comments']
        else:
            # This is the additional row (e.g., you could make it a blank row or a special one)
            row = table.add_row().cells
            row[0].text = "Additional Field"
            row[1].text = "Custom Description"
            row[2].text = "Comments for additional field"


    # Save the document
    filename = 'Generated_Document.docx'
    doc.save(filename)
    messagebox.showinfo("Success", f"Document saved as {filename}")

def on_generate():
    try:
        num_fields = int(num_fields_var.get())
        fields = []
        for i in range(num_fields):
            is_mandatory = mandatory_flags[i].get() == 'Yes'
            field = {
                'name': field_names[i].get(),
                'type': field_types[i].get(),
                'mandatory': is_mandatory,
                'description': generate_answer(field_types[i].get(), is_mandatory),
                'comments': additional_columns[i].get()
            }
            fields.append(field)

        generate_docx(fields)
    except Exception as e:
        messagebox.showerror("Error", str(e))

# def generate_description(field_type):
#     """Generate description based on field type"""
#     if field_type == "Text":
#         return "This is a mandatory field. User can enter text."
#     elif field_type == "Date":
#         return "This is a mandatory field. US format date is accepted in the field."
#     elif field_type == "Username":
#         return "This is a mandatory field. User can select a username from the people picker search box."
#     else:
#         return "No description available."

def create_fields():
    # Clear existing widgets
    for widget in fields_frame.winfo_children():
        widget.destroy()

    # Remove the previous "Generate Document" button if it exists
    for widget in button_frame.winfo_children():
        widget.destroy()

    try:
        num_fields = int(num_fields_var.get())
        if num_fields <= 0:
            raise ValueError("Number of fields must be positive.")
    except ValueError as e:
        messagebox.showerror("Error", str(e))
        return

    # Reset field variables
    global field_names, field_types, mandatory_flags, additional_columns
    field_names = []
    field_types = []
    mandatory_flags = []
    additional_columns = []

    # Create entry widgets for each field
    for i in range(num_fields):
        # Field name (first row)
        ttk.Label(fields_frame, text=f"Field Name {i + 1}:").grid(row=i * 2, column=0, padx=5, pady=5, sticky="w")
        name_var = tk.StringVar()
        field_names.append(name_var)
        ttk.Entry(fields_frame, textvariable=name_var, width=50).grid(row=i * 2, column=1, padx=5, pady=5, sticky="we")

        # Fields Type, Mandatory, Additional Comments (second row)
        ttk.Label(fields_frame, text="Data Type:").grid(row=i * 2 + 1, column=0, padx=5, pady=5, sticky="w")
        type_var = tk.StringVar()
        field_types.append(type_var)
        ttk.Combobox(fields_frame, textvariable=type_var, values=["Text", "Date", "Username", "Email"]).grid(row=i * 2 + 1, column=1, padx=5, pady=5, sticky="we")

        ttk.Label(fields_frame, text="Mandatory:").grid(row=i * 2 + 1, column=2, padx=5, pady=5, sticky="w")
        mandatory_var = tk.StringVar(value="No")
        mandatory_flags.append(mandatory_var)
        ttk.Combobox(fields_frame, textvariable=mandatory_var, values=["Yes", "No"]).grid(row=i * 2 + 1, column=3, padx=5, pady=5, sticky="we")

        ttk.Label(fields_frame, text="Additional Details:").grid(row=i * 2 + 1, column=4, padx=5, pady=5, sticky="w")
        additional_var = tk.StringVar()
        additional_columns.append(additional_var)
        ttk.Entry(fields_frame, textvariable=additional_var, width=50).grid(row=i * 2 + 1, column=5, padx=5, pady=5, sticky="we")

    # Add the "Generate Document" button
    ttk.Button(button_frame, text="Generate Document", command=on_generate).pack(pady=10)


# OPEN_AI_URL = 'https://api.openai.com/v1/chat/completions'
# PROXY_FOR_OPENAI = 'http://RyyRW8:e2Z6KQ@186.65.123.202:8000'
# OPENAI_API_KEY = ''

def generate_answer(field_type, mandatory):    
    
    proxy = {
        'http': PROXY_FOR_OPENAI,
        'https': PROXY_FOR_OPENAI
    }
    
    prompt = f"Generate a concise and user-friendly description for a field of type '{field_type}'. " \
             f"Keep the response brief and clear, focusing on how this field should be filled."

    payload = {
        "model": "gpt-3.5-turbo",
        "temperature": 0.5,
        "messages": [
            {
                "role": "user",
                "content": prompt
            }
        ]
    
    }
    
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {OPENAI_API_KEY}"
    }

    try:
        response = requests.post(OPEN_AI_URL, headers=headers, json=payload, proxies=proxy)
        response.raise_for_status()
        result = response.json()

        ai_description = result["choices"][0]["message"]["content"]

        if mandatory:
            ai_description = f"This is a mandatory field. {ai_description}"
        
        return ai_description

    except requests.exceptions.RequestException as e:
        return f"Request failed with status code {response.status_code}: {response.text}"


# GUI setup
root = tk.Tk()
root.title("Requirement Document Generator")

ttk.Label(root, text="Number of Fields:").pack(pady=5)
num_fields_var = tk.StringVar()
ttk.Entry(root, textvariable=num_fields_var).pack(pady=5)
ttk.Button(root, text="Create Fields", command=create_fields).pack(pady=5)

fields_frame = ttk.Frame(root)
fields_frame.pack(pady=10)

button_frame = ttk.Frame(root)
button_frame.pack(pady=10)

# Global variables for fields
field_names = []
field_types = []
mandatory_flags = []
additional_columns = []

root.mainloop()
